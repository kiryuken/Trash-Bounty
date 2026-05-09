from __future__ import annotations

from datetime import datetime
from io import BytesIO

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Mm, Pt, RGBColor


PERIOD_LABELS = {
	"weekly": "Mingguan (7 Hari)",
	"monthly": "Bulanan (30 Hari)",
	"alltime": "Semua Waktu",
}


def build_report(stats: dict, narrative: str, period: str) -> bytes:
	"""Build a cleanup impact report as a DOCX file and return the bytes."""
	doc = Document()
	section = doc.sections[0]
	section.page_width = Mm(210)
	section.page_height = Mm(297)
	section.left_margin = Inches(1)
	section.right_margin = Inches(1)
	section.top_margin = Inches(1)
	section.bottom_margin = Inches(1)

	title = doc.add_heading("Laporan Dampak Lingkungan TrashBounty Lumi", level=0)
	title.alignment = WD_ALIGN_PARAGRAPH.CENTER

	subtitle = doc.add_paragraph(
		f"Periode: {PERIOD_LABELS.get(period, period)} | Dibuat: {datetime.now().strftime('%d %B %Y')}"
	)
	subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
	if subtitle.runs:
		subtitle.runs[0].font.size = Pt(11)
		subtitle.runs[0].font.color.rgb = RGBColor(0x55, 0x55, 0x55)

	doc.add_paragraph()

	doc.add_heading("1. Ringkasan Eksekutif", level=1)
	doc.add_paragraph(narrative)

	doc.add_heading("2. Statistik Dampak", level=1)
	metrics_table = doc.add_table(rows=1, cols=2)
	metrics_table.style = "Table Grid"
	metrics_header = metrics_table.rows[0].cells
	metrics_header[0].text = "Metrik"
	metrics_header[1].text = "Nilai"
	for cell in metrics_header:
		for paragraph in cell.paragraphs:
			for run in paragraph.runs:
				run.bold = True

	rows = [
		("Total Bounty Diselesaikan", str(stats.get("total_completed", 0))),
		("Estimasi Berat Sampah Dibersihkan", f"{float(stats.get('total_weight_kg', 0)):.1f} kg"),
		("Total Poin Didistribusikan", f"{int(stats.get('total_points_awarded', 0)):,} poin"),
		("Total Reward Dibagikan", f"Rp {float(stats.get('total_reward_idr', 0)):,.0f}"),
	]
	for label, value in rows:
		cells = metrics_table.add_row().cells
		cells[0].text = label
		cells[1].text = value

	doc.add_paragraph()

	doc.add_heading("3. Jenis Sampah yang Dibersihkan", level=1)
	waste_types = stats.get("waste_types", []) or []
	if waste_types:
		waste_table = doc.add_table(rows=1, cols=3)
		waste_table.style = "Table Grid"
		waste_header = waste_table.rows[0].cells
		for index, label in enumerate(("Jenis Sampah", "Jumlah Bounty", "Rata-rata Keparahan")):
			waste_header[index].text = label
			for paragraph in waste_header[index].paragraphs:
				for run in paragraph.runs:
					run.bold = True

		for item in waste_types:
			cells = waste_table.add_row().cells
			cells[0].text = str(item.get("waste_type", "-"))
			cells[1].text = str(item.get("count", 0))
			cells[2].text = f"{float(item.get('avg_severity', 0)):.1f}/10"
	else:
		doc.add_paragraph("Belum ada data jenis sampah untuk periode ini.")

	footer = section.footer
	footer_paragraph = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
	footer_paragraph.text = f"Dibuat oleh TrashBounty Lumi | {datetime.now().strftime('%d/%m/%Y %H:%M')}"
	footer_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
	if footer_paragraph.runs:
		footer_paragraph.runs[0].font.size = Pt(9)

	buffer = BytesIO()
	doc.save(buffer)
	return buffer.getvalue()